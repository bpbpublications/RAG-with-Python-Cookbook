# document_paragraph_splitting.py
# This code demonstrates how to split a .docx file into individual paragraphs
# and create Document objects for each paragraph using the langchain library.
from docx import Document
from langchain_core.documents import Document as LangChainDocument

# 1. Load the .docx file
# Replace 'RAG.docx' with the path to your .docx file.
doc = Document("RAG.docx")

# 2. Extract paragraphs from the document
# Each paragraph is treated as a separate chunk.
paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# 3. Create Document objects for each paragraph
# Each Document represents a paragraph with its content and metadata.
split_docs = [
    LangChainDocument(page_content=para, metadata={"source": "RAG.docx", "paragraph_num": idx + 1})
    for idx, para in enumerate(paragraphs)
]

# 4. Print the resulting Document objects
# Each Document represents a paragraph with its content and metadata indicating the source and paragraph number.
for i, d in enumerate(split_docs):
    print(f"\nParagraph {i+1}:\n{d.page_content}")
